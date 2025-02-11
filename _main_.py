import streamlit as st

from docx import Document
from io import BytesIO

def create_document(name, age, options):
    document = Document()
    document.add_heading("User Information", 0)
    document.add_paragraph(f"Name: {name}")
    document.add_paragraph(f"Age: {age}")
    document.add_paragraph("Options Selected:")
    for option in options:
        document.add_paragraph(option, style='List Bullet')
    
    # Save the document to a BytesIO stream
    doc_stream = BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

st.title("Dynamic Document Generator")

name = st.text_input("Enter your name:")
age = st.number_input("Enter your age:", min_value=0, max_value=120)
options = st.multiselect("Choose your options:", ["Option A", "Option B", "Option C"])



if st.button("Generate Document"):
    if name and age and options:
        doc_stream = create_document(name, age, options)
        st.download_button(
            label="Download Document",
            data=doc_stream,
            file_name="user_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please fill in all fields.")


